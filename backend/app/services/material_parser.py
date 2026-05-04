"""
app.services.material_parser
────────────────────────────

Extracts plain text from teacher-uploaded supporting material so that
Gemma generation can be anchored to the source content.

Supported formats:
    .txt / .md  — plain text decoded as UTF-8 (with latin-1 fallback)
    .pdf        — extracted via pypdf if installed; raises if missing
    .docx       — extracted via python-docx if installed; raises if missing

The parser is intentionally permissive: it returns whatever text it can
recover, truncates to a configurable character limit, and never blocks
generation if extraction fails — it just returns an empty string with a
warning. The endpoint layer is responsible for surfacing that warning to
the teacher.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from pathlib import PurePath

logger = logging.getLogger(__name__)

# Generation prompts read teacher_material verbatim, so cap the length to
# protect Gemma's context window. Roughly 6k chars ≈ 1.5k tokens.
MAX_MATERIAL_CHARS = 6000

_TEXT_SUFFIXES = {".txt", ".md", ".markdown"}
_PDF_SUFFIXES = {".pdf"}
_DOCX_SUFFIXES = {".docx"}


@dataclass
class ParsedMaterial:
    text: str
    source_filename: str
    truncated: bool
    warning: str = ""

    def is_empty(self) -> bool:
        return not self.text.strip()


def _decode_text_bytes(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="replace")


def _extract_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "Soporte PDF no instalado. Añade 'pypdf' a requirements.txt."
        ) from exc

    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:
            logger.warning("[material_parser] PDF page extraction error: %s", exc)
    return "\n".join(p.strip() for p in parts if p and p.strip())


def _extract_docx(raw: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "Soporte DOCX no instalado. Añade 'python-docx' a requirements.txt."
        ) from exc

    doc = Document(io.BytesIO(raw))
    parts: list[str] = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_material(filename: str, raw: bytes) -> ParsedMaterial:
    """Best-effort text extraction from an uploaded teacher material file."""
    suffix = PurePath(filename).suffix.lower()
    safe_name = PurePath(filename).name or "material"

    if not raw:
        return ParsedMaterial(text="", source_filename=safe_name, truncated=False,
                              warning="Archivo vacío.")

    try:
        if suffix in _TEXT_SUFFIXES or suffix == "":
            text = _decode_text_bytes(raw)
        elif suffix in _PDF_SUFFIXES:
            text = _extract_pdf(raw)
        elif suffix in _DOCX_SUFFIXES:
            text = _extract_docx(raw)
        else:
            return ParsedMaterial(
                text="",
                source_filename=safe_name,
                truncated=False,
                warning=f"Formato '{suffix}' no soportado. Usa .txt, .md, .pdf o .docx.",
            )
    except RuntimeError as exc:
        # Missing optional dependency — bubble up a friendly message
        return ParsedMaterial(text="", source_filename=safe_name, truncated=False,
                              warning=str(exc))
    except Exception as exc:
        logger.exception("[material_parser] extraction failed for %s", safe_name)
        return ParsedMaterial(text="", source_filename=safe_name, truncated=False,
                              warning=f"No se pudo leer el archivo: {exc}")

    text = (text or "").strip()
    truncated = len(text) > MAX_MATERIAL_CHARS
    if truncated:
        text = text[:MAX_MATERIAL_CHARS].rstrip()

    return ParsedMaterial(
        text=text,
        source_filename=safe_name,
        truncated=truncated,
        warning="" if text else "No se extrajo texto útil del archivo.",
    )


def parse_inline_material(text: str) -> ParsedMaterial:
    """Wrap teacher material pasted directly into the dashboard textarea."""
    raw = (text or "").strip()
    truncated = len(raw) > MAX_MATERIAL_CHARS
    if truncated:
        raw = raw[:MAX_MATERIAL_CHARS].rstrip()
    return ParsedMaterial(
        text=raw,
        source_filename="inline",
        truncated=truncated,
    )
